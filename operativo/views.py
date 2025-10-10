from datetime import date, datetime, timedelta
import json
import io
import os
from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse, JsonResponse
from django.urls import reverse
from django.utils import timezone
from reportlab.lib.pagesizes import letter
from operativo import models
from djangocrud import settings
from .forms import IntegranteForm, NotaForm
from .models import AcuerdoOperativo, Integrante, Nota
import io
from datetime import date
from django.http import HttpResponse
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.units import cm
from .models import Integrante, Nota, AcuerdoOperativo 
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from reportlab.platypus import Paragraph
from django.urls import reverse
from django.views.decorators.http import require_POST
from .models import AcuerdoOperativo, ComentarioAcuerdo
from docx import Document


def operativo_view(request):
    if request.method == "POST" and 'agregar_integrante' in request.POST:
        form = IntegranteForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('operativo_index')
    else:
        form = IntegranteForm()

    integrantes = Integrante.objects.all().order_by('area__nombre', 'nombre_completo')

    query = request.GET.get('q')
    if query:
        integrantes = integrantes.filter(nombre_completo__icontains=query)
    seleccionados = []

    return render(request, "operativo/base.html", {
        'integrantes': integrantes,
        'form': form,
        'seleccionados': seleccionados,
        'fecha_actual': timezone.now(),
    })

# --- Logica para crear e insertar Notas  -------------------------------------------------------
def editar_nota(request, nota_id):
    nota = get_object_or_404(Nota, id=nota_id)
    if request.method == "POST":
        form = NotaForm(request.POST, instance=nota)
        if form.is_valid():
            form.save()
            return redirect('operativo_index')
    else:
        form = NotaForm(instance=nota)
    return render(request, 'operativo/partials/desarrollo.html', {'form': form, 'nota': nota})


def lista_notas(request):
    notas = Nota.objects.all().order_by('-fecha_creacion')
    return render(request, 'operativo/partials/desarrollo.html', {'notas': notas})


def guardar_todo(request):
    if request.method == "POST":
        data = json.loads(request.body)
        notas = data.get('notas', [])
        for n in notas:
            Nota.objects.create(apartado=n['apartado'], texto=n['texto'])
        return JsonResponse({'status': 'ok'})
    return JsonResponse({'status': 'error'}, status=400)


# ---Apartado para toda la logica de reglas -------------------------------------------------------------
def crear_acuerdo_operativo(request):
    integrantes = Integrante.objects.all().order_by('area__nombre', 'nombre_completo')
    unidades = list(range(1, 10))

    if request.method == "POST":
        filas = request.POST.getlist('numerador')
        unidades_post = request.POST.getlist('unidad')
        acuerdos = request.POST.getlist('acuerdo')
        unidad_paradas = request.POST.getlist('unidad_parada')
        fechas_limite = request.POST.getlist('fecha_limite')
        pendientes = request.POST.getlist('pendiente')
        responsables = request.POST.getlist('responsable')
        avances = request.POST.getlist('porcentaje_avance')

        for i in range(len(filas)):
            AcuerdoOperativo.objects.create(
                numerador=int(filas[i]),
                unidad=int(unidades_post[i]),
                acuerdo=acuerdos[i],
                unidad_parada=(unidad_paradas[i] == 'on') if i < len(unidad_paradas) else False,
                fecha_limite=fechas_limite[i],
                pendiente=(pendientes[i] == 'on') if i < len(pendientes) else True,
                responsable_id=int(responsables[i]),
                porcentaje_avance=int(avances[i])
            )
        return redirect('crear_acuerdo_operativo')

    return render(request, 'modulo/crear_acuerdo_operativo.html', {
        'integrantes': integrantes,
        'fecha_actual': timezone.now(),
        'unidades': unidades
    })


def historial_acuerdo_operativo(request):
    acuerdos = AcuerdoOperativo.objects.all().order_by('-fecha_creacion')
    query = request.GET.get('q')
    if query:
        acuerdos = acuerdos.filter(
            models.Q(acuerdo__icontains=query) |
            models.Q(numerador__icontains=query) |
            models.Q(responsable__nombre_completo__icontains=query)
        )
    return render(request, 'modulo/historial_acuerdo_operativo.html', {
        'acuerdos': acuerdos,
        'query': query,
    })

@require_POST
def editar_acuerdo_operativo(request, acuerdo_id):
    acuerdo = get_object_or_404(AcuerdoOperativo, id=acuerdo_id)
    
    acuerdo.unidad_parada = request.POST.get('unidad_parada') == 'on'
    acuerdo.acuerdo = request.POST.get('acuerdo')
    acuerdo.porcentaje_avance = int(request.POST.get('porcentaje_avance', acuerdo.porcentaje_avance))
    acuerdo.save()

    comentario_texto = request.POST.get('comentario', '')
    comentario_obj, created = ComentarioAcuerdo.objects.get_or_create(acuerdo=acuerdo)
    comentario_obj.texto = comentario_texto
    comentario_obj.save()

    return redirect('historial_acuerdo_operativo')  # ✅ usar el name de la URL


# --- Descarga.html toda la logica para descargar pdf con formato -------------------------------------------------------


def descarga(request):
    # ------------------------------
    # Filtrado de búsqueda para integrantes
    # ------------------------------
    query = request.GET.get("q", "")
    if query:
        integrantes = Integrante.objects.filter(
            models.Q(nombre_completo__icontains=query) |
            models.Q(puesto__icontains=query) |
            models.Q(area__nombre__icontains=query)
        ).order_by("area__nombre", "nombre_completo")
    else:
        integrantes = Integrante.objects.all().order_by("area__nombre", "nombre_completo")

    # ------------------------------
    # Filtro de fechas para notas y acuerdos
    # ------------------------------
    fecha_filtro = request.GET.get("fecha_filtro", "ultimos")  # ultimos, hoy, 7dias, 30dias, todos
    ahora = datetime.now()

    # Querysets base
    notas_qs = Nota.objects.all()
    acuerdos_qs = AcuerdoOperativo.objects.all()

    if fecha_filtro == "hoy":
        notas_qs = notas_qs.filter(fecha_creacion__date=ahora.date())
        acuerdos_qs = acuerdos_qs.filter(fecha_creacion__date=ahora.date())
    elif fecha_filtro == "7dias":
        fecha_inicio = ahora - timedelta(days=7)
        notas_qs = notas_qs.filter(fecha_creacion__gte=fecha_inicio)
        acuerdos_qs = acuerdos_qs.filter(fecha_creacion__gte=fecha_inicio)
    elif fecha_filtro == "30dias":
        fecha_inicio = ahora - timedelta(days=30)
        notas_qs = notas_qs.filter(fecha_creacion__gte=fecha_inicio)
        acuerdos_qs = acuerdos_qs.filter(fecha_creacion__gte=fecha_inicio)
    elif fecha_filtro == "todos":
        pass  # usamos todos los registros
    else:  # ultimos agregados por defecto
        notas_qs = notas_qs.order_by('-fecha_creacion')[:10]
        acuerdos_qs = acuerdos_qs.order_by('-fecha_creacion')[:10]

    # Ordenar si no usamos slicing
    if fecha_filtro in ["hoy", "7dias", "30dias", "todos"]:
        notas_qs = notas_qs.order_by('-fecha_creacion')
        acuerdos_qs = acuerdos_qs.order_by('-fecha_creacion')

    # ------------------------------
    # Integrantes seleccionados desde sesión (opcional)
    # ------------------------------
    seleccionados_ids = request.session.get("integrantes_seleccionados", [])
    seleccionados = Integrante.objects.filter(id__in=seleccionados_ids)

    form = IntegranteForm()

    return render(request, "modulo/descarga.html", {
        "integrantes": integrantes,
        "seleccionados": seleccionados,
        "notas": notas_qs,
        "acuerdos": acuerdos_qs,
        "form": form,
        "request": request,
        "fecha_filtro": fecha_filtro
    })



def descargar_word(request):
    if request.method != "POST":
        return HttpResponse("Método no permitido", status=405)

    # Obtener IDs seleccionados desde el form
    integrantes_ids = request.POST.getlist("integrantes")
    notas_ids = request.POST.getlist("notas_seleccionadas")
    acuerdos_ids = request.POST.getlist("acuerdos_seleccionados")

    # Traer objetos de la DB de manera segura
    try:
        integrantes = Integrante.objects.filter(id__in=integrantes_ids) if integrantes_ids else []
        notas = Nota.objects.filter(id__in=notas_ids) if notas_ids else []
        acuerdos = AcuerdoOperativo.objects.filter(id__in=acuerdos_ids) if acuerdos_ids else []
    except Exception as e:
        return HttpResponse(f"Error al obtener datos: {str(e)}", status=500)

    # Ruta segura de la plantilla
    plantilla_path = os.path.join(settings.MEDIA_ROOT, 'plantillas', 'Operativa.docx')
    if not os.path.exists(plantilla_path):
        return HttpResponse("Plantilla no encontrada.", status=404)

    try:
        # Cargar plantilla
        doc = Document(plantilla_path)

        # Función para reemplazar texto simple en todos los runs (para acuerdos y fecha)
        def reemplazar_marcador(doc, marcador, texto):
            for p in doc.paragraphs:
                for run in p.runs:
                    if marcador in run.text:
                        run.text = run.text.replace(marcador, texto)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                if marcador in run.text:
                                    run.text = run.text.replace(marcador, texto)

        # Función para reemplazar marcador con múltiples líneas (para integrantes y notas)
        def reemplazar_marcador_parrafos(doc, marcador, lista_textos):
            for p in doc.paragraphs:
                if marcador in p.text:
                    p.text = ''
                    for texto in lista_textos:
                        p.add_run(texto).add_break()
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if marcador in p.text:
                                p.text = ''
                                for texto in lista_textos:
                                    p.add_run(texto).add_break()

        # Preparar listas de texto a insertar
        lista_integrantes = [f"- {i.nombre_completo} ({i.puesto})" for i in integrantes] or ["Sin integrantes seleccionados"]

        # Notas sin fechas
        lista_notas = [f"- {n.texto}" for n in notas] or ["Sin notas seleccionadas"]

        # Acuerdos: todos los campos, separados por comas
        lista_acuerdos = []
        if acuerdos:
            for a in acuerdos:
                fila = f"{a.unidad}, {a.acuerdo}, {a.unidad_parada}, " \
                       f"{a.fecha_limite.strftime('%d/%m/%Y') if a.fecha_limite else 'N/A'}, " \
                       f"{a.pendiente}, {a.responsable.nombre_completo}, {a.porcentaje_avance}%"
                lista_acuerdos.append(f"- {fila}")
        else:
            lista_acuerdos = ["Sin acuerdos seleccionados"]

        # Fecha actual
        texto_fecha = timezone.now().strftime("%d/%m/%Y")

        # Reemplazar marcadores en el documento
        reemplazar_marcador_parrafos(doc, '{{integrantes}}', lista_integrantes)
        reemplazar_marcador_parrafos(doc, '{{notas}}', lista_notas)
        reemplazar_marcador(doc, '{{acuerdos}}', "\n".join(lista_acuerdos))
        reemplazar_marcador(doc, '{{fecha}}', texto_fecha)

        # Preparar respuesta para descarga
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename=Minuta_Operativa.docx'
        doc.save(response)
        return response

    except Exception as e:
        return HttpResponse(f"Error al generar Word: {str(e)}", status=500)



