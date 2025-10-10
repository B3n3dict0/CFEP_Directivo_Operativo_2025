# Python estándar
import io
from datetime import datetime, timedelta
import os
import json
from docx import Document
# Django
from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse, JsonResponse
from django.utils import timezone
from django.conf import settings
from django.views.decorators.csrf import csrf_exempt

# ReportLab
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.lib.utils import ImageReader
from django.shortcuts import render, get_object_or_404, redirect
from django.views.decorators.http import require_POST
# Modelos
from .models import ComentarioDirectivo, NotaDirectivo, AcuerdoDirectivo, Integrante

# Formularios
from .forms import NotaDirectivoForm, IntegranteForm
from directivo import models



# ---------------- VIEWS ----------------

def directivo_view(request):
    # Agregar integrante
    if request.method == "POST" and 'agregar_integrante' in request.POST:
        form = IntegranteForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('directivo_index')
    else:
        form = IntegranteForm()

    integrantes = Integrante.objects.all().order_by('area__nombre', 'nombre_completo')

    # Búsqueda
    query = request.GET.get('q')
    if query:
        integrantes = integrantes.filter(nombre_completo__icontains=query)

    # Integrantes seleccionados en session
    seleccionados_ids = request.session.get("integrantes_seleccionados", [])
    seleccionados = Integrante.objects.filter(id__in=seleccionados_ids)

    return render(request, "directivo/base.html", {
        'integrantes': integrantes,
        'form': form,
        'seleccionados': seleccionados,
        'fecha_actual': timezone.now(),
    })


def lista_notas_directivo(request):
    notas = NotaDirectivo.objects.all().order_by('-fecha_creacion')
    return render(request, 'directivo/partials/desarrollo.html', {'notas': notas})


def editar_nota_directivo(request, nota_id):
    nota = get_object_or_404(NotaDirectivo, id=nota_id)
    if request.method == "POST":
        form = NotaDirectivoForm(request.POST, instance=nota)
        if form.is_valid():
            form.save()
            return redirect('directivo_index')
    else:
        form = NotaDirectivoForm(instance=nota)
    return render(request, 'directivo/partials/desarrollo.html', {'form': form, 'nota': nota})


def guardar_todo_directivo(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            notas = data.get('notas', [])
            for n in notas:
                if 'apartado' in n and 'texto' in n:
                    NotaDirectivo.objects.create(apartado=n['apartado'], texto=n['texto'])
            return JsonResponse({'status': 'ok'})
        except Exception as e:
            return JsonResponse({'status': 'error', 'msg': str(e)}, status=400)
    return JsonResponse({'status': 'error'}, status=400)


# ---------------- REGLAS ----------------

def crear_acuerdo_directivo(request):
    integrantes = Integrante.objects.all().order_by('area__nombre', 'nombre_completo')
    unidades = list(range(1, 10))

    if request.method == "POST":
        filas = request.POST.getlist('numerador')
        unidades_post = request.POST.getlist('unidad')
        acuerdos = request.POST.getlist('acuerdo')
        unidad_paradas = request.POST.getlist('unidad_parada')
        fechas_limite = request.POST.getlist('fecha_limite')
        pendientes = request.POST.getlist('pendiente')
        responsables = request.POST.getlist('responsable')
        avances = request.POST.getlist('porcentaje_avance')

        for i in range(len(filas)):
            AcuerdoDirectivo.objects.create(
                numerador=int(filas[i]),
                unidad=int(unidades_post[i]),
                acuerdo=acuerdos[i],
                unidad_parada=(unidad_paradas[i] == 'on') if i < len(unidad_paradas) else False,
                fecha_limite=fechas_limite[i],
                pendiente=(pendientes[i] == 'on') if i < len(pendientes) else True,
                responsable_id=int(responsables[i]),
                porcentaje_avance=int(avances[i])
            )
        return redirect('crear_acuerdo_directivo')

    return render(request, 'modulo/crear_acuerdo_directivo.html', {
        'integrantes': integrantes,
        'fecha_actual': timezone.now(),
        'unidades': unidades
    })


def historial_acuerdos(request):
    acuerdos = AcuerdoDirectivo.objects.all().order_by('-fecha_creacion')
    query = request.GET.get('q')
    if query:
        acuerdos = acuerdos.filter(
            models.Q(acuerdo__icontains=query) |
            models.Q(numerador__icontains=query) |
            models.Q(responsable__nombre_completo__icontains=query)
        )
    return render(request, 'modulo/historial_acuerdo_directivo.html', {
        'acuerdos': acuerdos,
        'query': query,
    })

@require_POST
def editar_acuerdo_directivo(request, id):
    acuerdo = get_object_or_404(AcuerdoDirectivo, id=id)
    
    acuerdo.unidad_parada = request.POST.get("unidad_parada") == "on"
    acuerdo.acuerdo = request.POST.get("acuerdo")
    acuerdo.porcentaje_avance = int(request.POST.get("porcentaje_avance", acuerdo.porcentaje_avance))
    acuerdo.save()

    # Guardar o crear comentario
    comentario_texto = request.POST.get("comentario", "")
    comentario_obj, created = ComentarioDirectivo.objects.get_or_create(acuerdo=acuerdo)
    comentario_obj.texto = comentario_texto
    comentario_obj.save()

    return redirect("directivo_historial_acuerdos")


# ---------------- Selección y Descarga ----------------

def seleccionar_integrantes_directivo(request):
    integrantes = Integrante.objects.all().order_by("area__nombre", "nombre_completo")

    if request.method == "POST":
        seleccionados_ids = request.POST.getlist("integrantes")
        request.session["integrantes_seleccionados"] = [str(i) for i in seleccionados_ids]
        return redirect("descarga_directiva")

    seleccionados_ids = request.session.get("integrantes_seleccionados", [])
    seleccionados = Integrante.objects.filter(id__in=seleccionados_ids)

    return render(request, "modulo/descarga_directivo.html", {
        "integrantes": integrantes,
        "seleccionados": seleccionados,
    })



def descarga_directiva(request):
    from datetime import datetime, timedelta

    # Filtrado de búsqueda para integrantes
    query = request.GET.get("q", "")
    if query:
        integrantes = Integrante.objects.filter(
            models.Q(nombre_completo__icontains=query) |
            models.Q(puesto__icontains=query) |
            models.Q(area__nombre__icontains=query)
        ).order_by("area__nombre", "nombre_completo")
    else:
        integrantes = Integrante.objects.all().order_by("area__nombre", "nombre_completo")

    # Filtro de fechas para notas y acuerdos
    fecha_filtro = request.GET.get("fecha_filtro", "ultimos")  # valores: ultimos, hoy, 7dias, 30dias, todos
    ahora = datetime.now()

    # Construir queryset base para notas y acuerdos
    notas_qs = NotaDirectivo.objects.all()
    acuerdos_qs = AcuerdoDirectivo.objects.all()

    if fecha_filtro == "hoy":
        notas_qs = notas_qs.filter(fecha_creacion__date=ahora.date())
        acuerdos_qs = acuerdos_qs.filter(fecha_creacion__date=ahora.date())
    elif fecha_filtro == "7dias":
        fecha_inicio = ahora - timedelta(days=7)
        notas_qs = notas_qs.filter(fecha_creacion__gte=fecha_inicio)
        acuerdos_qs = acuerdos_qs.filter(fecha_creacion__gte=fecha_inicio)
    elif fecha_filtro == "30dias":
        fecha_inicio = ahora - timedelta(days=30)
        notas_qs = notas_qs.filter(fecha_creacion__gte=fecha_inicio)
        acuerdos_qs = acuerdos_qs.filter(fecha_creacion__gte=fecha_inicio)
    elif fecha_filtro == "todos":
        pass  # usamos todos los registros
    else:  # ultimos agregados por defecto
        notas_qs = notas_qs.order_by('-fecha_creacion')[:10]
        acuerdos_qs = acuerdos_qs.order_by('-fecha_creacion')[:10]

    # SOLO ordenar si no hemos hecho slicing
    if fecha_filtro in ["hoy", "7dias", "30dias", "todos"]:
        notas_qs = notas_qs.order_by('-fecha_creacion')
        acuerdos_qs = acuerdos_qs.order_by('-fecha_creacion')

    seleccionados_ids = request.session.get("integrantes_seleccionados", [])
    seleccionados = Integrante.objects.filter(id__in=seleccionados_ids)
    form = IntegranteForm()

    return render(request, "modulo/descarga_directivo.html", {
        "integrantes": integrantes,
        "seleccionados": seleccionados,
        "notas": notas_qs,
        "acuerdos": acuerdos_qs,
        "form": form,
        "request": request,
        "fecha_filtro": fecha_filtro
    })



# ---------------- Funciones WORD ----------------
def descargar_word_directiva(request):
    if request.method != "POST":
        return HttpResponse("Método no permitido", status=405)

    # Obtener IDs seleccionados desde el formulario
    integrantes_ids = request.POST.getlist("integrantes")
    notas_ids = request.POST.getlist("notas_seleccionadas")
    acuerdos_ids = request.POST.getlist("acuerdos_seleccionados")

    try:
        integrantes = Integrante.objects.filter(id__in=integrantes_ids) if integrantes_ids else []
        notas = NotaDirectivo.objects.filter(id__in=notas_ids) if notas_ids else []
        acuerdos = AcuerdoDirectivo.objects.filter(id__in=acuerdos_ids).order_by("numerador") if acuerdos_ids else []
    except Exception as e:
        return HttpResponse(f"Error al obtener datos: {str(e)}", status=500)

    # Ruta de la plantilla
    plantilla_path = os.path.join(settings.MEDIA_ROOT, 'plantillas', 'Directiva.docx')
    if not os.path.exists(plantilla_path):
        return HttpResponse("Plantilla no encontrada.", status=404)

    try:
        doc = Document(plantilla_path)

        # --- Función robusta para reemplazar marcador en cualquier run ---
        def reemplazar_marcador(doc, marcador, texto):
            for p in doc.paragraphs:
                if marcador in "".join(run.text for run in p.runs):
                    nuevo_texto = "".join(run.text for run in p.runs).replace(marcador, texto)
                    for run in p.runs:
                        run.text = ""
                    p.add_run(nuevo_texto)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if marcador in "".join(run.text for run in p.runs):
                                nuevo_texto = "".join(run.text for run in p.runs).replace(marcador, texto)
                                for run in p.runs:
                                    run.text = ""
                                p.add_run(nuevo_texto)

        # --- Función para listas (integrantes y notas) ---
        def reemplazar_marcador_parrafos(doc, marcador, lista_textos):
            for p in doc.paragraphs:
                if marcador in p.text:
                    p.text = ''
                    for texto in lista_textos:
                        p.add_run(texto).add_break()
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if marcador in p.text:
                                p.text = ''
                                for texto in lista_textos:
                                    p.add_run(texto).add_break()

        # Preparar listas de texto
        lista_integrantes = [f"- {i.nombre_completo} ({i.puesto})" for i in integrantes] or ["Sin integrantes seleccionados"]
        lista_notas = [f"- {n.texto}" for n in notas] or ["Sin notas seleccionadas"]

        # Reemplazar secciones
        reemplazar_marcador_parrafos(doc, '{{integrantes}}', lista_integrantes)
        reemplazar_marcador_parrafos(doc, '{{notas}}', lista_notas)

        # Fecha actual
        fecha_actual = timezone.now().strftime("%d/%m/%Y")
        reemplazar_marcador(doc, '{{fecha}}', fecha_actual)

        # --- Rellenar tabla de acuerdos ---
        tabla_obj = None
        fila_inicio = None
        for t in doc.tables:
            for idx, row in enumerate(t.rows):
                if any('{{acuerdos}}' in cell.text for cell in row.cells):
                    tabla_obj = t
                    fila_inicio = idx
                    break
            if tabla_obj:
                break

        if tabla_obj and fila_inicio is not None:
            max_rows = len(tabla_obj.rows)
            filas_disponibles = max_rows - fila_inicio

            for i, a in enumerate(acuerdos):
                if i >= filas_disponibles:
                    break
                target_row_index = fila_inicio + i
                cells = tabla_obj.rows[target_row_index].cells

                if len(cells) >= 4:
                    cells[0].text = str(a.numerador) if a.numerador is not None else str(i + 1)
                    cells[1].text = a.acuerdo or ""
                    cells[2].text = (a.responsable.nombre_completo if a.responsable else "-")
                    cells[3].text = a.fecha_limite.strftime('%d/%m/%Y') if a.fecha_limite else "-"
                else:
                    for c in cells:
                        c.text = ""

            # Limpiar filas sobrantes
            llenadas = min(len(acuerdos), filas_disponibles)
            for j in range(fila_inicio + llenadas, fila_inicio + filas_disponibles):
                if j >= max_rows:
                    break
                cells = tabla_obj.rows[j].cells
                for c in cells:
                    c.text = ""

        # --- Crear carpeta de respaldo y guardar copia ---
        backup_base = os.path.join(settings.BASE_DIR, 'respaldo_word_directivo')
        mes_actual = timezone.now().strftime("%Y-%m")
        backup_folder = os.path.join(backup_base, mes_actual)
        os.makedirs(backup_folder, exist_ok=True)

        nombre_archivo_usuario = f"Minuta_Directiva_{timezone.now().strftime('%Y%m%d_%H%M%S')}.docx"
        backup_path = os.path.join(backup_folder, nombre_archivo_usuario)
        doc.save(backup_path)

        # --- Generar respuesta para descarga ---
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename={nombre_archivo_usuario}'
        doc.save(response)
        return response

    except Exception as e:
        return HttpResponse(f"Error al generar Word: {str(e)}", status=500)
